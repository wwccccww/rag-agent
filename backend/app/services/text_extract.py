import io
import logging
import re
from pathlib import Path

import httpx
from pypdf import PdfReader


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            t = page.extract_text() or ""
            if t.strip():
                parts.append(f"\n--- Page {i + 1} ---\n{t}")
        return "\n".join(parts).strip()

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument  # lazy import
            doc = DocxDocument(io.BytesIO(data))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # 同时提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            result = "\n\n".join(paragraphs)
            logging.info("[Extract] docx → %d chars, %d paragraphs", len(result), len(paragraphs))
            return result
        except ImportError as e:
            raise RuntimeError("python-docx 未安装，请运行 pip install python-docx") from e

    if suffix == ".xlsx":
        try:
            import openpyxl  # lazy import
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts: list[str] = []
            for sheet in wb.worksheets:
                parts.append(f"## Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(" | ".join(cells))
            result = "\n".join(parts)
            logging.info("[Extract] xlsx → %d chars, %d sheets", len(result), len(wb.worksheets))
            return result
        except ImportError as e:
            raise RuntimeError("openpyxl 未安装，请运行 pip install openpyxl") from e

    # 默认按 UTF-8 文本处理（.txt / .md / .csv 等）
    return data.decode("utf-8", errors="replace")


def fetch_url(url: str, timeout: float = 15.0) -> tuple[str, str]:
    """抓取网页内容，返回 (纯文本, 页面标题)。依赖 beautifulsoup4 + lxml。"""
    try:
        from bs4 import BeautifulSoup  # lazy import，未安装时给出友好报错
    except ImportError as e:
        raise RuntimeError("beautifulsoup4 未安装，请运行 pip install beautifulsoup4 lxml") from e

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0 Safari/537.36"
        )
    }
    with httpx.Client(timeout=timeout, follow_redirects=True) as client:
        resp = client.get(url, headers=headers)
        resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "lxml")

    # 移除无用标签
    for tag in soup(["script", "style", "nav", "header", "footer", "aside", "form", "iframe"]):
        tag.decompose()

    title = (soup.title.string or "").strip() if soup.title else ""

    # 提取正文：优先 <article>/<main>，否则 <body>
    body = soup.find("article") or soup.find("main") or soup.body
    if body is None:
        body = soup

    lines = [line.strip() for line in body.get_text(separator="\n").splitlines() if line.strip()]
    text = "\n".join(lines)
    logging.info("[URL] fetched %s → %d chars (title: %s)", url, len(text), title)
    return text, title


_MD_HEADING = re.compile(r"^#{1,6}\s+\S")
_CLOSING_FENCE_RE = re.compile(r"^`{3,}\s*$")

# Parent-Child 分块：父块最小内容字符阈值（可通过 config 覆盖）
_PARENT_CHILD_MIN_PARENT_CHARS_DEFAULT = 200
_PARENT_CHILD_MAX_PARENT_CHARS_DEFAULT = 1500


def _heading_line_level(line: str) -> int | None:
    """与 _MD_HEADING 一致：行首 1–6 个 # 且其后有空白与正文。"""
    st = line.strip()
    if not _MD_HEADING.match(st):
        return None
    n = 0
    for c in st:
        if c == "#":
            n += 1
        else:
            break
    if n < 1 or n > 6:
        return None
    return n


def _is_closing_fence_line(line: str) -> bool:
    """围栏结束行：整行（去首尾空白）仅由至少 3 个 ` 与可选空白组成。"""
    return bool(_CLOSING_FENCE_RE.match(line.strip()))


def _iter_text_and_fence_spans(section: str) -> list[tuple[str, str]]:
    """将一节正文拆成交替的 prose 与 fenced code（含首尾 ``` 行），围栏内原文不改。"""
    lines = section.replace("\r\n", "\n").split("\n")
    out: list[tuple[str, str]] = []
    text_buf: list[str] = []
    i = 0
    n = len(lines)

    def flush_text() -> None:
        if not text_buf:
            return
        block = "\n".join(text_buf)
        text_buf.clear()
        if block.strip():
            out.append(("text", block))

    while i < n:
        line = lines[i]
        if line.strip().startswith("```"):
            flush_text()
            fence_lines = [line]
            i += 1
            while i < n:
                fence_lines.append(lines[i])
                if len(fence_lines) > 1 and _is_closing_fence_line(lines[i]):
                    i += 1
                    break
                i += 1
            joined = "\n".join(fence_lines)
            if joined.strip():
                out.append(("fence", joined))
            continue
        text_buf.append(line)
        i += 1
    flush_text()
    return out


def _section_to_mixed_units(section: str) -> list[tuple[str, str]]:
    """text 单元按空行拆成段落；fence 单元整块保留。"""
    units: list[tuple[str, str]] = []
    for kind, block in _iter_text_and_fence_spans(section):
        if kind == "text":
            for p in block.split("\n\n"):
                ps = p.strip()
                if ps:
                    units.append(("text", ps))
        elif block.strip():
            units.append(("fence", block))
    return units


def _first_heading_in_section(section: str) -> str | None:
    """节内第一个 Markdown 标题行（去掉 #），供纯代码块 meta 回落。"""
    for line in section.replace("\r\n", "\n").split("\n")[:48]:
        st = line.strip()
        if _MD_HEADING.match(st):
            return st.lstrip("#").strip()[:160]
    return None


def _merge_short_intro_before_fence_units(
    units: list[tuple[str, str]], max_intro: int
) -> list[tuple[str, str]]:
    """将「连续短正文 + 紧随的围栏」合并为一个 fence 单元（多段 text 会先拼成引言，再与 ``` 块合并）。"""
    if max_intro <= 0 or len(units) < 2:
        return units
    out: list[tuple[str, str]] = []
    i = 0
    while i < len(units):
        if units[i][0] == "text":
            j = i
            texts: list[str] = []
            while j < len(units) and units[j][0] == "text":
                texts.append(units[j][1].strip())
                j += 1
            if j < len(units) and units[j][0] == "fence":
                intro = "\n\n".join(t for t in texts if t)
                fence = units[j][1].strip()
                if intro and "```" not in intro and len(intro) <= max_intro:
                    out.append(("fence", intro + "\n\n" + fence))
                    i = j + 1
                    continue
            for t in texts:
                if t:
                    out.append(("text", t))
            i = j
            continue
        out.append(units[i])
        i += 1
    return out


def _heading_from_chunk_start(content: str, max_lines: int = 8) -> str | None:
    for ln in content.split("\n")[:max_lines]:
        st = ln.strip()
        if _MD_HEADING.match(st):
            return st.lstrip("#").strip()[:120]
    return None


def _short_section_label_for_continuation(raw: str | None, max_chars: int) -> str | None:
    """超长围栏续块前缀用：去掉 #、取首行、截断。"""
    if not raw or not str(raw).strip():
        return None
    s = str(raw).strip().split("\n", 1)[0].strip()
    while s.startswith("#"):
        s = s.lstrip("#").strip()
    s = s.strip()
    if not s:
        return None
    if len(s) <= max_chars:
        return s
    return s[: max(1, max_chars - 1)] + "…"


def _looks_like_markdown(text: str, filename: str | None) -> bool:
    fn = (filename or "").lower()
    if fn.endswith(".md"):
        return True
    t = text.lstrip()
    if t.startswith("#"):
        return True
    return "\n## " in text or "\n### " in text


def _split_oversized_paragraph(p: str, max_chars: int, overlap: int) -> list[str]:
    """长段按句号/换行优先切，避免硬截断把语义切碎。"""
    if len(p) <= max_chars:
        return [p] if p.strip() else []
    out: list[str] = []
    start = 0
    n = len(p)
    while start < n:
        end = min(start + max_chars, n)
        if end < n:
            ws = max(start, end - 140)
            window = p[ws:end]
            cut = -1
            for sep in ("。\n", "。", "！\n", "！", "？\n", "？", ". ", ".\n", "\n\n", "\n", "；", "; "):
                idx = window.rfind(sep)
                if idx != -1:
                    cut = ws + idx + len(sep)
                    break
            if cut > start:
                end = cut
        piece = p[start:end].strip()
        if piece:
            out.append(piece)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return out


def _split_oversized_fence(
    p: str,
    max_chars: int,
    _overlap: int,
    *,
    continuation_label: str | None = None,
) -> list[str]:
    """超长围栏块仅在换行处切分，避免在代码/XML 行内硬截断。

    子块之间**不使用**与正文相同的字符 _overlap：回退 overlap 容易落在行中，产生半截标签。
    若提供 continuation_label 且切出多块，从第 2 块起在正文前加 ``[节：… · 续]`` 便于向量与阅读关联小节。
    """
    if len(p) <= max_chars:
        return [p] if p.strip() else []
    out: list[str] = []
    start = 0
    n = len(p)
    while start < n:
        end = min(start + max_chars, n)
        if end < n:
            ws = max(start, end - 400)
            window = p[ws:end]
            cut = window.rfind("\n")
            if cut != -1:
                end = ws + cut + 1
            else:
                segment = p[start:end]
                cut2 = segment.rfind("\n")
                if cut2 != -1:
                    end = start + cut2 + 1
        piece = p[start:end].strip()
        if piece:
            out.append(piece)
        if end >= n:
            break
        start = end
    if continuation_label and len(out) > 1:
        pref = f"[节：{continuation_label} · 续]"
        merged: list[str] = [out[0]]
        for piece in out[1:]:
            head = piece.lstrip()[:40]
            if head.startswith("[节：") and "· 续]" in head:
                merged.append(piece)
            else:
                merged.append(f"{pref}\n\n{piece}".strip())
        return merged
    return out


def _pack_paragraphs(paragraphs: list[str], max_chars: int, overlap: int) -> list[str]:
    chunks: list[str] = []
    buf = ""
    for p in paragraphs:
        if len(buf) + len(p) + 2 <= max_chars:
            buf = f"{buf}\n\n{p}" if buf else p
        else:
            if buf:
                chunks.append(buf)
            if len(p) <= max_chars:
                buf = p
            else:
                for piece in _split_oversized_paragraph(p, max_chars, overlap):
                    chunks.append(piece)
                buf = ""
        while len(buf) > max_chars:
            chunks.append(buf[:max_chars])
            buf = buf[max_chars - overlap :]
    if buf:
        chunks.append(buf)
    return [c.strip() for c in chunks if c.strip()]


def _pack_mixed_units(
    units: list[tuple[str, str]],
    max_chars: int,
    overlap: int,
    *,
    section_heading: str | None = None,
    fence_continuation_prefix: bool = True,
    continuation_title_max_chars: int = 72,
) -> list[str]:
    """在 text 单元上沿用段落合并与长段软切；fence 单元整块输出，超长时仅按换行切。"""
    chunks: list[str] = []
    buf = ""
    cont_label: str | None = None
    if fence_continuation_prefix and section_heading and section_heading.strip():
        cont_label = _short_section_label_for_continuation(
            section_heading, max(8, continuation_title_max_chars)
        )
    for kind, p in units:
        if kind == "fence":
            if buf and len(buf) + len(p) + 2 <= max_chars:
                buf = f"{buf}\n\n{p}"
                while len(buf) > max_chars:
                    chunks.append(buf[:max_chars])
                    buf = buf[max_chars - overlap :]
                continue
            if buf:
                chunks.append(buf)
                buf = ""
            if len(p) <= max_chars:
                chunks.append(p)
            else:
                chunks.extend(
                    _split_oversized_fence(
                        p, max_chars, overlap, continuation_label=cont_label
                    )
                )
            continue
        # text
        if len(buf) + len(p) + 2 <= max_chars:
            buf = f"{buf}\n\n{p}" if buf else p
        else:
            if buf:
                chunks.append(buf)
            if len(p) <= max_chars:
                buf = p
            else:
                for piece in _split_oversized_paragraph(p, max_chars, overlap):
                    chunks.append(piece)
                buf = ""
        while len(buf) > max_chars:
            chunks.append(buf[:max_chars])
            buf = buf[max_chars - overlap :]
    if buf:
        chunks.append(buf)
    return [c.strip() for c in chunks if c.strip()]


def _markdown_section_tuples(text: str) -> list[tuple[str, str | None]]:
    """按 Markdown 标题行分节；每节附带「标题面包屑」（父级 # 标题 / 当前节标题，不含 #）。

    用于 `#####` 等小节在向量与 meta 中仍能关联上级 `####` 接口段落。
    """
    text = text.replace("\r\n", "\n").strip()
    if not text:
        return []
    lines = text.split("\n")
    out: list[tuple[str, str | None]] = []
    buf: list[str] = []
    stack: list[tuple[int, str]] = []

    def crumb_str() -> str | None:
        if not stack:
            return None
        parts = [t[1] for t in stack if t[1]]
        return " / ".join(parts) if parts else None

    for line in lines:
        st = line.strip()
        is_h = bool(_MD_HEADING.match(st))
        lev = _heading_line_level(line) if is_h else None
        if is_h and buf:
            sec = "\n".join(buf).strip()
            if sec:
                out.append((sec, crumb_str()))
            if lev is not None:
                while stack and stack[-1][0] >= lev:
                    stack.pop()
                stack.append((lev, st[lev:].strip()[:200]))
            buf = [line]
        elif is_h and not buf:
            if lev is not None:
                while stack and stack[-1][0] >= lev:
                    stack.pop()
                stack.append((lev, st[lev:].strip()[:200]))
            buf.append(line)
        else:
            buf.append(line)
    if buf:
        sec = "\n".join(buf).strip()
        if sec:
            out.append((sec, crumb_str()))
    return out if out else [(text, None)]


def _markdown_sections(text: str) -> list[str]:
    """兼容：仅返回分节正文列表。"""
    return [s for s, _ in _markdown_section_tuples(text)]


# ── Parent-Child 分块 ──────────────────────────────────────────────────────────

def _section_heading_level(section_text: str) -> int | None:
    """从节首行提取标题级别（如 `### 5.1.2` → 3）。"""
    for line in section_text.splitlines():
        st = line.strip()
        if _MD_HEADING.match(st):
            return _heading_line_level(line)
    return None


def _detect_parent_level(section_tuples: list[tuple[str, str | None]]) -> int | None:
    """自动检测「父块层级」：叶子层（最深标题级别）的上一层。

    例如文档有 ###/####/##### 三层，叶子层 = 5，父块层 = 4（`####`）。
    若文档只有一层标题，返回该层（每个标题各成一组）。
    """
    levels: list[int] = []
    for sec, _ in section_tuples:
        lv = _section_heading_level(sec)
        if lv is not None:
            levels.append(lv)
    if not levels:
        return None
    leaf_level = max(levels)
    min_level = min(levels)
    # 父块层 = 叶子层 - 1；若叶子层就是最浅层则用叶子层（每节自成父块）
    return max(min_level, leaf_level - 1)


def _group_into_parent_sections(
    section_tuples: list[tuple[str, str | None]],
    min_parent_chars: int,
    max_parent_chars: int,
) -> list[list[tuple[str, str | None]]]:
    """把细粒度分节列表合并为「父节」组。

    核心策略：**按标题层级**分组，而非字符数。
    1. 先检测「父块层级」= 文档叶子层（最深标题）的上一层。
       例如 ###/####/##### 文档：叶子层 5 → 父块层 4（每个 `####` 开启一个新组）。
    2. 遇到标题层级 ≤ 父块层级 → 新开一个父块组（等于或更浅的标题是新的语义单元起点）。
    3. 更深的标题（子节）追加进当前组。
    4. 若合并后超过 max_parent_chars，按字符数二次切分（防止单父块过长）。
    这样 ##### 1.1.1 / 1.1.2 / 1.1.3 三个子节会合并进同一个以 #### 1.1注册 开头的父块中。
    """
    if not section_tuples:
        return []

    parent_lv = _detect_parent_level(section_tuples)

    # 若无法检测层级，回落到按字符数简单分组
    if parent_lv is None:
        return [[t] for t in section_tuples]

    # ── 按层级分组 ─────────────────────────────────────────────────
    raw_groups: list[list[tuple[str, str | None]]] = []
    cur_group: list[tuple[str, str | None]] = []

    for tup in section_tuples:
        sec, _ = tup
        lv = _section_heading_level(sec)
        # 遇到「父块层级或更浅」的标题 → 开新组
        if lv is not None and lv <= parent_lv:
            if cur_group:
                raw_groups.append(cur_group)
            cur_group = [tup]
        else:
            # 更深的子节 / 无标题内容 → 追加进当前组
            cur_group.append(tup)

    if cur_group:
        raw_groups.append(cur_group)

    # ── 二次切分：防止单父块超过 max_parent_chars ──────────────────
    result: list[list[tuple[str, str | None]]] = []
    for group in raw_groups:
        total = sum(len(s) for s, _ in group)
        if total <= max_parent_chars:
            result.append(group)
            continue
        # 超长：按字符数二次切分（保持至少一个 tup 一组）
        sub: list[tuple[str, str | None]] = []
        sub_chars = 0
        for tup in group:
            s_len = len(tup[0])
            if sub and sub_chars + s_len > max_parent_chars:
                result.append(sub)
                sub = [tup]
                sub_chars = s_len
            else:
                sub.append(tup)
                sub_chars += s_len
        if sub:
            result.append(sub)

    return result


def chunk_text_hierarchical(
    text: str,
    max_chars: int,
    overlap: int,
    *,
    filename: str | None = None,
    markdown_by_heading: bool = True,
    markdown_fence_aware: bool = True,
    merge_intro_before_fence_max_chars: int = 320,
    fence_continuation_prefix: bool = True,
    continuation_title_max_chars: int = 72,
    min_parent_chars: int = _PARENT_CHILD_MIN_PARENT_CHARS_DEFAULT,
    max_parent_chars: int = _PARENT_CHILD_MAX_PARENT_CHARS_DEFAULT,
) -> list[tuple[str, dict, list[tuple[str, dict]]]]:
    """Parent-Child 分块。

    返回 list of (parent_content, parent_meta, children)，其中：
    - parent_content / parent_meta：父块完整内容与元数据（用于喂给模型）
    - children：list of (child_content, child_meta)，子块（用于向量检索）

    若文档不是 Markdown 或不适合父子切分，则每个 chunk 作为自身的「单节父块」，
    children 为空列表（调用方统一按普通 chunk 入库，is_index_chunk=True）。

    父块与子块共享 section_heading；父块额外有 meta["is_parent"]=True。
    """
    text = text.strip()
    if not text:
        return []

    use_md = markdown_by_heading and _looks_like_markdown(text, filename)
    if not use_md:
        # 非 Markdown：直接退化为普通切块，每块自成父块
        plain = chunk_text(
            text, max_chars, overlap,
            filename=filename,
            markdown_by_heading=False,
            markdown_fence_aware=False,
        )
        return [(c, m, []) for c, m in plain]

    raw_sections = _markdown_section_tuples(text)
    groups = _group_into_parent_sections(raw_sections, min_parent_chars, max_parent_chars)

    result: list[tuple[str, dict, list[tuple[str, dict]]]] = []
    global_idx = 0  # 父块编号（用于 chunk_index）

    for group in groups:
        # ── 父块内容：所有组内节拼接 ──────────────────────
        parent_content = "\n\n".join(sec for sec, _ in group).strip()
        # 父块面包屑：取组内首节的面包屑（最能代表父块语义）
        first_crumb = group[0][1]
        parent_heading = (
            (first_crumb.strip() if first_crumb and first_crumb.strip() else None)
            or _first_heading_in_section(group[0][0])
        )
        parent_meta: dict = {
            "chunk_index": global_idx,
            "is_parent": True,
        }
        if parent_heading:
            parent_meta["section_heading"] = parent_heading[:200]

        if len(group) == 1:
            # 单节组：可能自身较大，需要进一步切块；子块即是切出来的块
            sec, sec_crumb = group[0]
            sec_heading = (
                (sec_crumb.strip() if sec_crumb and sec_crumb.strip() else None)
                or _first_heading_in_section(sec)
            )
            children_raw = _cut_single_section(
                sec, sec_heading, max_chars, overlap,
                markdown_fence_aware=markdown_fence_aware,
                merge_intro_before_fence_max_chars=merge_intro_before_fence_max_chars,
                fence_continuation_prefix=fence_continuation_prefix,
                continuation_title_max_chars=continuation_title_max_chars,
                start_idx=global_idx,
            )
            if len(children_raw) <= 1:
                # 单节且只切出一块 → 退化：只有一个块，无须父子分离
                result.append((parent_content, parent_meta, []))
                global_idx += 1
                continue
            children = [(c, m) for c, m in children_raw]
        else:
            # 多节合并组：每个小节切块作为子块
            children: list[tuple[str, dict]] = []
            child_idx = global_idx
            for sec, sec_crumb in group:
                sec_heading = (
                    (sec_crumb.strip() if sec_crumb and sec_crumb.strip() else None)
                    or _first_heading_in_section(sec)
                )
                for c, m in _cut_single_section(
                    sec, sec_heading, max_chars, overlap,
                    markdown_fence_aware=markdown_fence_aware,
                    merge_intro_before_fence_max_chars=merge_intro_before_fence_max_chars,
                    fence_continuation_prefix=fence_continuation_prefix,
                    continuation_title_max_chars=continuation_title_max_chars,
                    start_idx=child_idx,
                ):
                    children.append((c, m))
                    child_idx += 1

        result.append((parent_content, parent_meta, children))
        global_idx += 1 + len(children)

    return result


def _cut_single_section(
    sec: str,
    sec_heading: str | None,
    max_chars: int,
    overlap: int,
    *,
    markdown_fence_aware: bool,
    merge_intro_before_fence_max_chars: int,
    fence_continuation_prefix: bool,
    continuation_title_max_chars: int,
    start_idx: int,
) -> list[tuple[str, dict]]:
    """对单个节正文做内部切块，返回 (content, meta) 列表。"""
    if markdown_fence_aware:
        units = _section_to_mixed_units(sec)
        if merge_intro_before_fence_max_chars > 0:
            units = _merge_short_intro_before_fence_units(units, merge_intro_before_fence_max_chars)
        raw_chunks = _pack_mixed_units(
            units, max_chars, overlap,
            section_heading=sec_heading,
            fence_continuation_prefix=fence_continuation_prefix,
            continuation_title_max_chars=continuation_title_max_chars,
        ) if units else ([sec.strip()] if sec.strip() else [])
    else:
        paragraphs = [p.strip() for p in sec.split("\n\n") if p.strip()] or ([sec] if sec.strip() else [])
        raw_chunks = _pack_paragraphs(paragraphs, max_chars, overlap)

    out: list[tuple[str, dict]] = []
    for i, c in enumerate(raw_chunks):
        if not c.strip():
            continue
        m: dict = {"chunk_index": start_idx + i}
        h = sec_heading or _heading_from_chunk_start(c)
        if h:
            m["section_heading"] = h[:200]
        out.append((c.strip(), m))
    return out


def chunk_text(
    text: str,
    max_chars: int,
    overlap: int,
    *,
    filename: str | None = None,
    markdown_by_heading: bool = True,
    markdown_fence_aware: bool = True,
    merge_intro_before_fence_max_chars: int = 320,
    fence_continuation_prefix: bool = True,
    continuation_title_max_chars: int = 72,
) -> list[tuple[str, dict]]:
    """
    将全文切块。Markdown（.md 或含 ## 标题）在 markdown_by_heading 为 True 时先按标题分节；
    markdown_fence_aware 且判定为 Markdown 时，节内识别 ``` 围栏，围栏内不按句号/短窗切分，
    超长围栏仅在换行处切；可选将不超过 merge_intro_before_fence_max_chars 的紧邻引言并入围栏单元，
    避免「例如：」单独成块。超长围栏多块时可选为续块加 ``[节：… · 续]`` 前缀（fence_continuation_prefix）。
    每节 `meta.section_heading` 优先为**标题面包屑**（如 ``父级小节 / 当前 #####``）；无栈信息时回落为节内首个 `#` 标题。
    """
    text = text.strip()
    if not text:
        return []

    use_md = markdown_by_heading and _looks_like_markdown(text, filename)
    if use_md:
        raw_section_items = _markdown_section_tuples(text)
    else:
        raw_section_items = [(text, None)]

    pieces: list[tuple[str, str | None]] = []
    for sec, sec_breadcrumb in raw_section_items:
        sec_heading = (sec_breadcrumb.strip() if sec_breadcrumb and sec_breadcrumb.strip() else None) or _first_heading_in_section(sec)
        if use_md and markdown_fence_aware:
            units = _section_to_mixed_units(sec)
            if merge_intro_before_fence_max_chars > 0:
                units = _merge_short_intro_before_fence_units(units, merge_intro_before_fence_max_chars)
            if units:
                for c in _pack_mixed_units(
                    units,
                    max_chars,
                    overlap,
                    section_heading=sec_heading,
                    fence_continuation_prefix=fence_continuation_prefix,
                    continuation_title_max_chars=continuation_title_max_chars,
                ):
                    pieces.append((c, sec_heading))
            elif sec.strip():
                for c in _pack_paragraphs([sec.strip()], max_chars, overlap):
                    pieces.append((c, sec_heading))
        else:
            paragraphs = [p.strip() for p in sec.split("\n\n") if p.strip()]
            if not paragraphs:
                paragraphs = [sec] if sec.strip() else []
            for c in _pack_paragraphs(paragraphs, max_chars, overlap):
                pieces.append((c, sec_heading))

    out: list[tuple[str, dict]] = []
    for idx, (c, sec_fallback) in enumerate(pieces):
        page = None
        if "--- Page " in c:
            try:
                line = c.split("\n", 1)[0]
                if line.startswith("--- Page ") and "---" in line:
                    num = line.replace("--- Page ", "").split(" ", 1)[0]
                    page = int(num)
            except (ValueError, IndexError):
                page = None
        fb = sec_fallback.strip()[:200] if sec_fallback and sec_fallback.strip() else None
        # 分节已给出面包屑时优先写入 meta，避免块首仅 ##### 时丢失父级 ####
        h = fb or _heading_from_chunk_start(c)
        meta: dict = {"chunk_index": idx, "page": page}
        if h:
            meta["section_heading"] = h[:200]
        out.append((c.strip(), meta))
    return out
